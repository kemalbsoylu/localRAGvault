import io

import docx
import pypdf
import pytest

from core.extractors import extract_text_from_file


def test_extract_txt_utf8() -> None:
    content = b"Clean UTF-8 text string."
    result = extract_text_from_file(content, "sample.txt")
    assert result == "Clean UTF-8 text string."


def test_extract_txt_latin1_fallback() -> None:
    content = b"Latin-1 specific: \xe9\xe8\xe0"
    result = extract_text_from_file(content, "legacy.txt")
    assert "Latin-1 specific:" in result


def test_extract_csv_formatting() -> None:
    csv_bytes = b"host,port,status\nlocalhost,8000,active\ndb-server,5432,standby"
    result = extract_text_from_file(csv_bytes, "servers.csv")
    assert "[Row 1] host: localhost, port: 8000, status: active." in result
    assert "[Row 2] host: db-server, port: 5432, status: standby." in result


def test_extract_csv_empty_raises() -> None:
    with pytest.raises(ValueError, match="empty or missing headers"):
        extract_text_from_file(b"", "empty.csv")


def test_extract_json_dict_flattening() -> None:
    json_bytes = b'{"database": {"engine": "postgres", "version": 16}, "active": true}'
    result = extract_text_from_file(json_bytes, "config.json")
    assert "database.engine: postgres" in result
    assert "database.version: 16" in result
    assert "active: True" in result


def test_extract_json_list_formatting() -> None:
    json_bytes = b'[{"name": "Alice", "role": "admin"}, {"name": "Bob", "role": "dev"}]'
    result = extract_text_from_file(json_bytes, "users.json")
    assert "[Record 1] name: Alice, role: admin" in result
    assert "[Record 2] name: Bob, role: dev" in result


def test_extract_docx_in_memory() -> None:
    """Verify DOCX paragraphs and structured tables convert cleanly into searchable text streams."""
    doc = docx.Document()
    doc.add_paragraph("Introductory paragraph text inside DOCX.")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Header Alpha"
    table.rows[0].cells[1].text = "Header Beta"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = extract_text_from_file(buf.getvalue(), "report.docx")
    assert "Introductory paragraph text inside DOCX." in result
    assert "Header Alpha | Header Beta" in result


def test_extract_pdf_empty_or_scanned_raises() -> None:
    """Verify extracting text from a blank/scanned PDF raises a descriptive ValueError."""
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=72, height=72)

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)

    with pytest.raises(ValueError, match="No extractable text found in PDF"):
        extract_text_from_file(buf.getvalue(), "scanned_image.pdf")


def test_unsupported_extension_raises() -> None:
    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text_from_file(b"some binary data", "archive.zip")
