import csv
import io
import json
from typing import Any, List

import docx
import pypdf

from core.logging_config import logger


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Dispatches file bytes to the appropriate format-specific parser.
    Returns cleaned UTF-8 text ready for chunking.
    """
    if (
        not file_bytes
        or len(file_bytes.strip() if isinstance(file_bytes, bytes) else file_bytes) == 0
    ):
        raise ValueError(f"File '{filename}' is empty.")

    lower_name = filename.lower()
    try:
        if lower_name.endswith((".txt", ".md")):
            return _extract_txt(file_bytes, filename)
        elif lower_name.endswith(".pdf"):
            return _extract_pdf(file_bytes, filename)
        elif lower_name.endswith(".docx"):
            return _extract_docx(file_bytes, filename)
        elif lower_name.endswith(".csv"):
            return _extract_csv(file_bytes, filename)
        elif lower_name.endswith(".json"):
            return _extract_json(file_bytes, filename)
        else:
            raise ValueError(
                f"Unsupported file format: '{filename}'. Supported extensions: .txt, .md, .pdf, .docx, .csv, .json"
            )
    except ValueError as ve:
        logger.warning(f"Extraction validation failed for '{filename}': {ve}")
        raise
    except Exception as e:
        logger.error(f"Unexpected extraction failure for file '{filename}': {e}")
        raise


def _extract_txt(file_bytes: bytes, filename: str) -> str:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning(
            f"UTF-8 decode failed for '{filename}', attempting fallback latin-1 decoding..."
        )
        try:
            text = file_bytes.decode("latin-1")
        except Exception as exc:
            raise ValueError(f"File '{filename}' could not be decoded as text.") from exc

    if not text.strip():
        raise ValueError(f"File '{filename}' is empty.")
    return text


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"PDF file '{filename}' is corrupted or unreadable.") from e

    page_texts: List[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        cleaned_text = text.strip()
        if cleaned_text:
            page_texts.append(f"--- [Page {i + 1}] ---\n{cleaned_text}")

    if not page_texts:
        raise ValueError(
            f"No extractable text found in PDF '{filename}'. It may be empty or a scanned image."
        )
    return "\n\n".join(page_texts)


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX file '{filename}' is corrupted or unreadable.") from e

    content_parts: List[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            content_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                content_parts.append(" | ".join(row_data))

    if not content_parts:
        raise ValueError(f"No extractable text found in DOCX '{filename}' (file is empty).")
    return "\n\n".join(content_parts)


def _extract_csv(file_bytes: bytes, filename: str) -> str:
    decoded_text = file_bytes.decode("utf-8", errors="replace").strip()
    if not decoded_text:
        raise ValueError(f"CSV file '{filename}' is empty.")

    text_stream = io.StringIO(decoded_text)
    reader = csv.DictReader(text_stream)

    if not reader.fieldnames or all(not fn.strip() for fn in reader.fieldnames):
        raise ValueError(f"CSV file '{filename}' appears to be empty or missing headers.")

    sentences: List[str] = []
    for row_idx, row in enumerate(reader, start=1):
        row_statements = [
            f"{key.strip()}: {val.strip()}"
            for key, val in row.items()
            if key and val and val.strip()
        ]
        if row_statements:
            sentences.append(f"[Row {row_idx}] " + ", ".join(row_statements) + ".")

    if not sentences:
        raise ValueError(f"No valid data rows found in CSV '{filename}'.")
    return "\n".join(sentences)


def _extract_json(file_bytes: bytes, filename: str) -> str:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError(f"JSON file '{filename}' is empty.")

    try:
        data = json.loads(text)
    except json.JSONDecodeError as err:
        raise ValueError(f"Invalid JSON format in file '{filename}': {err.msg}") from err

    if data is None or data == {} or data == []:
        raise ValueError(f"JSON file '{filename}' contains empty data structure.")

    def _flatten_to_sentences(obj: Any, prefix: str = "") -> List[str]:
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                key_label = f"{prefix}.{k}" if prefix else str(k)
                if isinstance(v, (dict, list)):
                    lines.extend(_flatten_to_sentences(v, key_label))
                elif v is not None and str(v).strip():
                    lines.append(f"{key_label}: {v}")
        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                item_label = f"{prefix}[{idx}]"
                if isinstance(item, (dict, list)):
                    lines.extend(_flatten_to_sentences(item, item_label))
                elif item is not None and str(item).strip():
                    lines.append(f"{item_label}: {item}")
        else:
            if obj is not None and str(obj).strip():
                lines.append(f"{prefix}: {obj}")
        return lines

    if isinstance(data, list):
        formatted_items = []
        for idx, item in enumerate(data, start=1):
            if isinstance(item, dict):
                statements = [
                    f"{k}: {v}" for k, v in item.items() if v is not None and str(v).strip()
                ]
                if statements:
                    formatted_items.append(f"[Record {idx}] " + ", ".join(statements))
            elif item is not None and str(item).strip():
                formatted_items.append(f"[Record {idx}] {item}")

        if not formatted_items:
            raise ValueError(f"JSON list in file '{filename}' contains no elements.")
        return "\n".join(formatted_items)

    elif isinstance(data, dict):
        lines = _flatten_to_sentences(data)
        if not lines:
            raise ValueError(f"JSON object in file '{filename}' contains no valid fields.")
        return "\n".join(lines)
    else:
        return str(data)
